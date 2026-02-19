"""
Reads TestExecution_Rewrite.md and generates a Word document (TE_Generated.docx)
with one table per test execution record, matching the format of TE.docx.

Usage:
    python generate_te_docx.py

Output:
    docs/reports/TE_Generated.docx
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt

# The table style used in the original TE.docx
TABLE_STYLE = "Grid Table 4"


def parse_execution_records(md_path: str) -> list[dict]:
    """Parse the markdown file and return a list of test execution record dicts."""
    text = Path(md_path).read_text(encoding="utf-8")

    # Split on ### headers (each test execution record)
    sections = re.split(r"^### ", text, flags=re.MULTILINE)
    records = []

    for section in sections:
        if not section.strip():
            continue

        # Extract test case ID and name from header line
        header_match = re.match(r"(TC-\S+)\s*—\s*(.+)", section.split("\n")[0])
        if not header_match:
            continue

        test_case_id = header_match.group(1).strip()
        test_name = header_match.group(2).strip()

        # Extract metadata table fields
        testing_tools = ""
        testing_type = ""
        for m in re.finditer(
            r"\|\s*\*\*Testing Tools Used\*\*\s*\|\s*(.+?)\s*\|", section
        ):
            testing_tools = m.group(1).strip()
        for m in re.finditer(
            r"\|\s*\*\*Testing Type\*\*\s*\|\s*(.+?)\s*\|", section
        ):
            testing_type = m.group(1).strip()

        # Extract execution steps
        steps = []
        steps_block = re.search(
            r"\*\*Execution Steps:\*\*\s*\n((?:\d+\..+\n?)+)", section
        )
        if steps_block:
            for line in steps_block.group(1).strip().split("\n"):
                step_match = re.match(r"\d+\.\s*(.+)", line.strip())
                if step_match:
                    steps.append(step_match.group(1).strip())

        # Extract execution record rows (the table after "Test Execution Records:")
        exec_rows = []
        table_lines = re.findall(
            r"^\|\s*(\d+)\s*\|(.+)\|$", section, re.MULTILINE
        )
        for num, rest in table_lines:
            cols = [c.strip() for c in rest.split("|")]
            if len(cols) >= 6:
                exec_rows.append(
                    {
                        "num": num.strip(),
                        "tester": cols[0],
                        "date": cols[1],
                        "result": cols[2],
                        "status": cols[3],
                        "defect": cols[4],
                        "correction": cols[5],
                    }
                )

        # Extract execution summary
        summary = ""
        summary_match = re.search(
            r"\*\*Execution Summary:\*\*\s*(.+?)(?:\n\n|\n---|\Z)",
            section,
            re.DOTALL,
        )
        if summary_match:
            summary = summary_match.group(1).strip()

        records.append(
            {
                "test_case_id": test_case_id,
                "test_name": test_name,
                "testing_tools": testing_tools,
                "testing_type": testing_type,
                "steps": steps,
                "exec_rows": exec_rows,
                "summary": summary,
            }
        )

    return records


def build_table(doc: Document, rec: dict):
    """Create one TE table in the document matching the TE.docx format."""
    n_steps = len(rec["steps"])
    n_exec = len(rec["exec_rows"])
    total_rows = 4 + n_steps + 1 + 1 + n_exec + 1
    COLS = 8

    table = doc.add_table(rows=total_rows, cols=COLS, style=TABLE_STYLE)

    def set_cell(row_idx, col_idx, text, font_size=Pt(10)):
        cell = table.rows[row_idx].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = font_size

    def merge_and_set(row_idx, col_start, col_end, text, font_size=Pt(10)):
        cell = table.rows[row_idx].cells[col_start].merge(
            table.rows[row_idx].cells[col_end]
        )
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = font_size

    row = 0

    # --- Metadata rows (0-3): label in cols 0-1, value in cols 2-7 ---
    metadata = [
        ("Project Name", "AIBI CV for Manufacturing"),
        ("Test Case ID", rec["test_case_id"]),
        ("Testing Tools Used", rec["testing_tools"]),
        ("Testing Type", rec["testing_type"]),
    ]
    for label, value in metadata:
        merge_and_set(row, 0, 1, label)
        merge_and_set(row, 2, 7, value)
        row += 1

    # --- Execution Steps rows ---
    for i, step_text in enumerate(rec["steps"]):
        merge_and_set(row, 0, 1, "Execution Steps:")
        set_cell(row, 2, str(i + 1))
        merge_and_set(row, 3, 7, step_text)
        row += 1

    # --- "Test Execution Records:" divider row ---
    merge_and_set(row, 0, 7, "Test Execution Records:")
    row += 1

    # --- Execution records header row ---
    set_cell(row, 0, "#")
    merge_and_set(row, 1, 2, "Tester")
    set_cell(row, 3, "Test Date")
    set_cell(row, 4, "Actual Result")
    set_cell(row, 5, "Status")
    set_cell(row, 6, "Defect")
    set_cell(row, 7, "Correction")
    row += 1

    # --- Execution record data rows ---
    for er in rec["exec_rows"]:
        set_cell(row, 0, er["num"])
        merge_and_set(row, 1, 2, er["tester"])
        set_cell(row, 3, er["date"])
        set_cell(row, 4, er["result"])
        set_cell(row, 5, er["status"])
        defect = er["defect"] if er["defect"] not in ("\u2014", "-", "\u2013") else "-"
        correction = (
            er["correction"] if er["correction"] not in ("\u2014", "-", "\u2013") else "-"
        )
        set_cell(row, 6, defect)
        set_cell(row, 7, correction)
        row += 1

    # --- Execution Summary row ---
    merge_and_set(row, 0, 1, "Execution Summary")
    merge_and_set(row, 2, 7, rec["summary"])


def main():
    script_dir = Path(__file__).parent
    md_path = script_dir / "TestExecution_Rewrite.md"
    template_path = script_dir / "TE.docx"
    out_path = script_dir / "TE_Generated.docx"

    print(f"Reading: {md_path}")
    records = parse_execution_records(str(md_path))
    print(f"Parsed {len(records)} test execution records")

    # Use TE.docx as template to inherit the "Grid Table 4" style
    doc = Document(str(template_path))

    # Clear the template content (remove existing table and any paragraphs)
    for table in doc.tables:
        table._element.getparent().remove(table._element)
    # Remove all body paragraphs except keep one empty one
    body = doc.element.body
    for p in list(body.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")):
        body.remove(p)

    # Set default font
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    for i, rec in enumerate(records):
        if i > 0:
            doc.add_page_break()
        build_table(doc, rec)

    doc.save(str(out_path))
    print(f"Generated: {out_path}")


if __name__ == "__main__":
    main()
